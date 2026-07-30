import sys
import subprocess

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

doc = Document()

# Set standard margins (1 inch all around)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Helper function to add styled headings
def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(16, 78, 139) # Deep Blue
    p.paragraph_format.space_after = Pt(4)

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(70, 130, 180) # Steel Blue
    p.paragraph_format.space_after = Pt(24)

def add_heading_1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(16, 78, 139)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)

def add_heading_2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(40, 100, 150)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

def add_body(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Calibri'
        r_bold.font.size = Pt(11)
        r_bold.font.bold = True
        r_bold.font.color.rgb = RGBColor(30, 30, 30)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(40, 40, 40)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Calibri'
        r_bold.font.size = Pt(11)
        r_bold.font.bold = True
        r_bold.font.color.rgb = RGBColor(30, 30, 30)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(40, 40, 40)
    return p

def add_callout(text, title="NOTE FOR MENTOR EXPLANATION"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    
    # Background color light blue/grey
    shading_xml = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4F8"/>')
    cell._tc.get_or_add_tcPr().append(shading_xml)
    
    # Left border thick blue
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="36" w:space="0" w:color="104E8B"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.1)
    
    r_title = p.add_run(f"📌 {title}\n")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(11)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(16, 78, 139)
    
    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(10.5)
    r_text.font.italic = True
    r_text.font.color.rgb = RGBColor(50, 50, 50)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# Document Content Generation

add_title("MediNow — Project Documentation")
add_subtitle("Simple & Clear Explanation of Mobile App, Web App, AI Features, and Backend Technologies")

add_callout(
    "This document explains everything about the MediNow project in very simple English. "
    "You can easily present this to your mentor or professor to explain what the project is, why it was created, how it works, and what technologies were used to build it.",
    "QUICK OVERVIEW FOR MENTORS"
)

# SECTION 1
add_heading_1("1. What is MediNow?")
add_body(
    "MediNow is a complete digital healthcare platform designed to help patients manage their daily medicines, "
    "read doctor prescriptions automatically using Artificial Intelligence (AI), track their pill schedules, "
    "and connect with nearby pharmacies to order medicines easily."
)
add_body(
    "The project consists of two user applications (a Mobile App for smartphones and a Web App for web browsers) "
    "powered by a single backend server."
)

add_heading_2("Project Summary at a Glance:")
add_bullet("MediNow (Healthcare & Pharmacy Assistant)", "Project Name: ")
add_bullet("Smart AI Healthcare, Medicine Management, & Local Pharmacy Delivery", "Main Concept: ")
add_bullet("Patients, Elderly People, Chronic Illness Patients, and Pharmacy Store Owners", "Target Users: ")
add_bullet("Flutter (Mobile App), React.js (Web App), FastAPI Python (Backend Server), SQLite (Database), and Google Gemini AI.", "Main Technologies: ")

# SECTION 2
add_heading_1("2. Why Did We Build MediNow? (The Problem & Solution)")
add_heading_2("The Problem:")
add_bullet("Doctor prescription handwriting is often very difficult to read for ordinary people.")
add_bullet("Patients frequently forget to take their daily medicines on time.")
add_bullet("Medicines run out suddenly or expire without the patient noticing.")
add_bullet("Finding nearby open pharmacies that have specific medicines in stock takes time and effort.")

add_heading_2("Our Solution (MediNow):")
add_bullet("AI Prescription Reader: Automatically scans doctor prescription photos and reads medicine names and dosages.")
add_bullet("Pill Alarm & Tracker: Sends reminders to take pills on time and keeps a daily log of taken vs. skipped medicines.")
add_bullet("Stock & Expiry Alerts: Warns users when medicine supply drops below 5 days or is close to expiring.")
add_bullet("GPS Pharmacy & Hospital Finder: Shows real nearby pharmacies on a map and allows ordering medicines directly.")

# SECTION 3
add_heading_1("3. Core Features (What Can MediNow Do?)")

add_heading_2("Feature 1: AI Prescription Scanner (OCR)")
add_body("Users can take a photo of a doctor's handwritten prescription paper. Google Gemini AI scans the photo, deciphers the handwriting, and automatically extracts:")
add_bullet("Medicine names (e.g., Paracetamol, Dolo 650)")
add_bullet("Dosage strength (e.g., 650 mg, 500 mg)")
add_bullet("Schedule & frequency (e.g., Twice daily - morning and night)")
add_bullet("Instructions (e.g., Take after meals)")
add_body("Users can add these extracted medicines straight into their pill schedule with just one click.")

add_heading_2("Feature 2: Smart Pill Reminders & Daily Tracker")
add_body("Acts as a personal alarm clock for daily medicines. Patients can log if they took or skipped a dose. The app saves this history to show compliance visual graphs.")

add_heading_2("Feature 3: Low-Stock Warning & Expiry Guard")
add_bullet("Smart Refill: Calculates how many pills are left and reminds the user when they have less than 5 days of medicine remaining.")
add_bullet("Expiry Warning: Highlights any medicine in the home inventory expiring within 30 days so users don't take expired pills.")

add_heading_2("Feature 4: AI Health Assistant (Chatbot)")
add_body("A friendly medical AI chatbot powered by Gemini AI. Patients can ask health questions, symptom advice, or details about their medicines. The AI knows what medicines the patient is taking and provides safe, context-aware answers.")

add_heading_2("Feature 5: GPS Pharmacy & Hospital Finder")
add_body("Uses real GPS coordinates to find nearby pharmacies and emergency hospitals within a user-chosen radius (e.g., 5 km). It shows exact distance, address, phone number, and whether the store is currently open.")

add_heading_2("Feature 6: Online Medicine Ordering & Live Delivery Tracking")
add_body("Patients can search medicines, compare prices, add items to a cart, and place an order. The app tracks the order stage step-by-step (Pending -> Confirmed -> Out for Delivery -> Delivered).")

add_heading_2("Feature 7: Adherence Analytics & Risk Score")
add_body("Calculates a 30-day compliance percentage score (e.g., 85% Adherence). It labels the user risk as Low Risk, Medium Risk, or High Risk, and provides simple AI tips on how to build better medicine habits.")

add_heading_2("Feature 8: Clinical Doctor Summary Export")
add_body("Generates a simple, exportable summary report of all recent medicines, missed doses, and logged symptoms that patients can show to their doctor during appointments.")

# SECTION 4
add_heading_1("4. User Portals (Who Uses MediNow?)")
add_body("MediNow has two distinct user portals depending on who is using the system:")

add_heading_2("A. Patient Portal (App & Web)")
add_bullet("Scan prescriptions and view digital records.")
add_bullet("Set daily medicine alarms and log taken/skipped doses.")
add_bullet("Check medicine inventory and reorder low-stock medicines.")
add_bullet("Chat with AI Doctor assistant.")
add_bullet("Locate nearby open pharmacies and hospitals.")
add_bullet("Track live medicine delivery status.")

add_heading_2("B. Pharmacy Store Owner Portal (Web App)")
add_bullet("Dashboard showing daily sales, pending orders, and earnings.")
add_bullet("Accept or reject customer medicine orders.")
add_bullet("Assign delivery drivers with phone numbers.")
add_bullet("Manage stock quantities and set prices for medicines.")

# SECTION 5
add_heading_1("5. Tools & Technologies Used (The Tech Stack)")
add_body("Here is a simple, easy-to-explain table of all tools, programming languages, and frameworks used in MediNow:")

# Table setup
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

# Header Row
hdr_cells = table.rows[0].cells
headers = ["Layer", "Technology / Tool", "Language", "What It Does (Simple Role)"]
widths = [Inches(1.2), Inches(1.8), Inches(1.0), Inches(2.5)]

shading_hdr = parse_xml(f'<w:shd {nsdecls("w")} w:fill="104E8B"/>')
for i, h in enumerate(headers):
    hdr_cells[i].width = widths[i]
    p = hdr_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(h)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)

# Data Rows
data = [
    ("Mobile Frontend", "Flutter", "Dart", "Cross-platform mobile app for Android & iOS smartphones."),
    ("Web Frontend", "React.js + Vite", "JavaScript / HTML / CSS", "Fast web application with modern Glassmorphism UI styling."),
    ("Backend API", "FastAPI + Uvicorn", "Python 3.11", "Central web server that processes requests and connects database & AI."),
    ("Database", "SQLite & SQLAlchemy", "SQL / Python", "Stores user profiles, medicine stock, orders, and dosage logs."),
    ("AI Vision OCR", "Google Gemini 2.0 Flash", "Python API", "Reads doctor handwritten prescription images and extracts medicine text."),
    ("AI Chatbot", "Google Gemini 2.0 Flash", "Python API", "Smart medical assistant that answers patient questions."),
    ("GPS Map API", "OpenStreetMap (Overpass)", "REST API", "Finds real nearby pharmacies and hospitals based on live GPS location."),
    ("Push Alerts", "Firebase Cloud Messaging", "Python / Dart", "Sends pill reminder notifications directly to user phones."),
    ("Security", "JWT Tokens & Bcrypt", "Python", "Encrypts passwords and keeps user login sessions secure.")
]

for row_idx, (layer, tech, lang, desc) in enumerate(data):
    row_cells = table.add_row().cells
    bg_color = "F9FBFD" if row_idx % 2 == 0 else "FFFFFF"
    
    for c_idx, val in enumerate([layer, tech, lang, desc]):
        row_cells[c_idx].width = widths[c_idx]
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
        row_cells[c_idx]._tc.get_or_add_tcPr().append(shd)
        
        p = row_cells[c_idx].paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(val)
        run.font.name = 'Calibri'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(40, 40, 40)
        if c_idx == 1:
            run.font.bold = True

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# SECTION 6
add_heading_1("6. How MediNow Works Step-by-Step (Workflows)")

add_heading_2("Workflow A: Scanning a Prescription")
add_bullet("Patient takes a picture of a handwritten doctor prescription paper in the app.")
add_bullet("The app sends the picture to our FastAPI server.")
add_bullet("FastAPI passes the picture to Google Gemini 2.0 Flash Vision AI.")
add_bullet("Gemini AI reads the handwriting and returns clean JSON data (Medicine name, dosage, frequency).")
add_bullet("The user verifies the detected medicines and taps 'Save to My Schedule'.")

add_heading_2("Workflow B: Ordering Medicine from a Local Pharmacy")
add_bullet("Patient searches for a medicine or views low-stock items.")
add_bullet("App fetches nearby pharmacies via OpenStreetMap API using patient's GPS coordinates.")
add_bullet("Patient chooses a pharmacy, adds medicines to cart, and taps 'Place Order'.")
add_bullet("The order appears instantly on the Pharmacy Owner's Web Dashboard.")
add_bullet("Pharmacy owner approves the order and assigns a delivery driver.")
add_bullet("Patient sees real-time status update: Pending -> Confirmed -> Out for Delivery -> Delivered.")

# SECTION 7
add_heading_1("7. Database Structure (Simple Data Models)")
add_body("MediNow uses 10 simple database tables to store all information cleanly:")
add_bullet("Users: Stores user profiles (Name, Email, Encrypted Password, Role as Patient or Pharmacy).")
add_bullet("Pharmacies: Stores pharmacy details (Name, Address, GPS Latitude/Longitude, Rating).")
add_bullet("Medicines: Master list of all available medicines.")
add_bullet("PharmacyStock: Keeps track of medicine quantity and price at each pharmacy.")
add_bullet("UserMedicines: Patient's current home inventory, daily dosage, and remaining pills.")
add_bullet("Prescriptions: Uploaded prescription image links and AI-extracted medicine data.")
add_bullet("Reminders: Configured alarm times for taking pills.")
add_bullet("DoseLogs: History of every dose taken or skipped by the patient.")
add_bullet("Orders: Online medicine orders with total price, delivery address, and status.")
add_bullet("HealthLogs: Patient logged symptoms, side effects, and severity.")

# SECTION 8
add_heading_1("8. How to Run the Project (Commands Summary)")
add_bullet("Run Backend Server: cd MediNow/backend && uvicorn main:app --reload")
add_bullet("Run Web App: cd MediNow/web && npm run dev")
add_bullet("Run Mobile App: cd MediNow/frontend && flutter run")

add_callout(
    "Summary for Mentor Speech:\n"
    "\"MediNow is an end-to-end AI healthcare ecosystem built with Flutter for mobile, React for web, "
    "and FastAPI in Python for backend. It uses Google Gemini AI to read messy doctor handwriting, "
    "reminds patients to take pills on time, warns before stock runs out, and uses OpenStreetMap GPS "
    "to connect patients directly with local pharmacies for medicine delivery.\"",
    "READY-TO-USE MENTOR PRESENTATION SCRIPT"
)

output_path = r"e:\Medicine\docs\MediNow_Project_Explanation.docx"
doc.save(output_path)
print(f"SUCCESS: Document generated at {output_path}")
